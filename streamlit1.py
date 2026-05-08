import tempfile
import streamlit as st
import os
from pathlib import Path
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    CSVLoader
)
from langchain_core.documents import Document

import pandas as pd
import pdfplumber

from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import Ollama

st.set_page_config(page_title="RAG-chatbot", layout="wide")
st.title("CHAT WITH YOUR FILE")
uploaded_file = st.sidebar.file_uploader(
    "Upload a file",
    type=["pdf", "docx", "csv"],
    key="uploader"
)

if "last_file" not in st.session_state:
    st.session_state.last_file = None

if "token_limit" not in st.session_state:
    st.session_state.token_limit = 500

current_file = uploaded_file.name if uploaded_file else None

if current_file != st.session_state.last_file:
    st.session_state.token_limit = 500   # reset value
    st.session_state.last_file = current_file

token_limit = st.sidebar.number_input(
    "Enter Output Token Limit",
    min_value=100,
    max_value=10000,
    value=st.session_state.token_limit,
    step=100
)

st.session_state.token_limit = token_limit

current_dir = os.path.dirname(os.path.abspath(__file__))
INDEX_PATH = os.path.join(current_dir, "faiss_index")
DATA_PATH = os.path.join(current_dir, "data")

os.makedirs(DATA_PATH, exist_ok=True)


@st.cache_resource
def load_resources():
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return embeddings

embeddings = load_resources()

# initialize session state
if "last_uploaded" not in st.session_state:
    st.session_state.last_uploaded = None
    if "messages" not in st.session_state:
        st.session_state.messages = []

# auto delete when uploader cleared
if uploaded_file is None and st.session_state.last_uploaded:

    old_path = os.path.join(
        DATA_PATH,
        st.session_state.last_uploaded
    )

    if os.path.exists(old_path):
        os.remove(old_path)

    if os.path.exists(INDEX_PATH):

        import shutil
        shutil.rmtree(INDEX_PATH)

    st.session_state.last_uploaded = None


with st.sidebar:

    if uploaded_file is not None:

        file_path = os.path.join(
            DATA_PATH,
            uploaded_file.name
        )

        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        temp_path = file_path

        st.session_state.last_uploaded = uploaded_file.name

        file_extension = os.path.splitext(
            uploaded_file.name
        )[1].lower()

        # PDF
        if file_extension == ".pdf":

            text = ""

            loader = PyPDFLoader(temp_path)

            pdf_docs = loader.load()

            for doc in pdf_docs:
                text += doc.page_content + "\n"

            with pdfplumber.open(temp_path) as pdf:

                for page in pdf.pages:

                    tables = page.extract_tables()

                    for table in tables:

                        for row in table:

                            cleaned_row = [
                                str(cell) if cell else ""
                                for cell in row
                            ]

                            text += " | ".join(cleaned_row) + "\n"

            documents = [Document(page_content=text)]

        # DOCX
        elif file_extension == ".docx":

            doc = DocxDocument(temp_path)

            text = ""

            for para in doc.paragraphs:
                text += para.text + "\n"

            for table in doc.tables:

                for row in table.rows:

                    row_text = [
                        cell.text for cell in row.cells
                    ]

                    text += " | ".join(row_text) + "\n"

            documents = [Document(page_content=text)]

        # CSV
        elif file_extension == ".csv":

            df = pd.read_csv(temp_path)

            text = df.to_string(index=False)

            documents = [Document(page_content=text)]

        else:

            st.error("Unsupported file type")
            st.stop()

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=600,
            chunk_overlap=100
        )

        chunks = splitter.split_documents(documents)

        vectorstore = FAISS.from_documents(
            chunks,
            embeddings
        )

        vectorstore.save_local(INDEX_PATH)

        st.success("Index Created!")


if os.path.exists(INDEX_PATH):
    
    vectorstore = FAISS.load_local(INDEX_PATH, embeddings, allow_dangerous_deserialization=True)
    for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
    
    query = st.chat_input("Ask a question...")

    if query:

        # show user message (RIGHT side)
        st.chat_message("user").markdown(query)
        st.session_state.messages.append({"role": "user", "content": query})

        docs = vectorstore.similarity_search(query, k=6)

        context = "\n\n".join([doc.page_content for doc in docs])

        prompt = f"""
        You are a strict question-answering system.

        RULES:
        - Use ONLY the information provided in the context.
        - Do NOT use any outside knowledge.
        - If the answer is not clearly present in the context, say: "I don't know based on the provided data."

        Context:
        {context}

        Question:
        {query}

        Answer:
        """

        llm = Ollama(model="phi3:mini", num_predict=token_limit)

        response_box = st.chat_message("assistant")
        placeholder = response_box.empty()

        full_response = ""

        for chunk in llm.stream(prompt):
            full_response += chunk
            placeholder.markdown(full_response)

        st.session_state.messages.append(
            {"role": "assistant", "content": full_response}
        )
else:
    st.info("upload a file to begin.")